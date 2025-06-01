from pathlib import Path
from docx import Document

# Directory paths for templates and output.
template_dir = Path("templates")
output_dir = Path("ready")
output_dir.mkdir(exist_ok=True)


def replace_placeholders(doc: Document, replacements: dict):
    # Replacing placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, val)
    # Replacing placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, val)

def process_templates(replacements: dict):
    for i, file in enumerate(template_dir.glob("*.docx"), start=1):
        try:
            doc = Document(file)
            replace_placeholders(doc, replacements)
            output_path = output_dir / f"ready_{i}.docx"
            doc.save(output_path)
        except Exception as e:
            print(f"Error processing {file}: {e}")